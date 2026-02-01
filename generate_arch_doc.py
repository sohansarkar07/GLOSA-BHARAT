from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """
    Set cell's border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "val": "single", "color": "#00FF00", "space": "0"},    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if not create
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('top', 'start', 'bottom', 'end'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key, value in edge_data.items():
                if key == 'color':
                    element.set(qn('w:color'), value)
                else:
                    element.set(qn('w:{}'.format(key)), value)

def create_detailed_architecture_doc():
    doc = Document()

    # Title Section
    title = doc.add_heading('SYSTEM ARCHITECTURE: GLOSA-BHARAT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Modern AI-Based Traffic Flow & Speed Advisory Ecosystem')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Theme: AI for Atmanirbhar Bharat')
    doc.add_paragraph('Focus: Full-Stack Microservices Architecture')

    # Layer 1
    doc.add_heading('LAYER 1: PERCEPTION (Edge Intelligence)', level=1)
    p1 = doc.add_paragraph()
    p1.add_run('Technology Stack: ').bold = True
    p1.add_run('Python, OpenCV, YOLOv8 (AI Objects Detection)')
    doc.add_paragraph(
        "This layer transforms raw visual data from urban CCTV cameras into structured telemetry. "
        "It uses indigenous AI models to identify vehicles, motorcycles, and emergency services in real-time, "
        "calculating queue density and lane occupancy without specialized foreign sensors.",
        style='List Bullet'
    )

    # Layer 2
    doc.add_heading('LAYER 2: INTEGRATION (The Nexus)', level=1)
    p2 = doc.add_paragraph()
    p2.add_run('Technology Stack: ').bold = True
    p2.add_run('Node.js, Express, Axios, REST API')
    doc.add_paragraph(
        "Acts as the communication gateway. It aggregates real-time JSON payloads from multiple junctions "
        "and handles authentication and secure data routing. It ensures that traffic telemetry is normalized "
        "before being processed by the Analytical Layer.",
        style='List Bullet'
    )

    # Layer 3
    doc.add_heading('LAYER 3: ANALYTICAL (The Brain)', level=1)
    p3 = doc.add_paragraph()
    p3.add_run('Technology Stack: ').bold = True
    p3.add_run('Proprietary GLOSA Algorithms, predictive Math Models')
    doc.add_paragraph(
        "The core logic where 'Green Wave' optimization occurs. The system calculates the intersection of "
        "current signal states (Red/Grey/Green) and incoming vehicle velocities to determine the "
        "Recommended Speed (GLOSA) for drivers.",
        style='List Bullet'
    )

    # Layer 4
    doc.add_heading('LAYER 4: INTERACTION (Digital Twin)', level=1)
    p4 = doc.add_paragraph()
    p4.add_run('Technology Stack: ').bold = True
    p4.add_run('React.js, Socket.io, Leaflet GIS, Framer Motion')
    doc.add_paragraph(
        "The visualization layer. It provides a real-time 'Control Room' interface for authorities and "
        "a minimalist advisory feed for drivers. Using WebSockets, it ensures zero-latency transmission "
        "of speed recommendations directly to the V2I (Vehicle-to-Infrastructure) interface.",
        style='List Bullet'
    )

    # Visual Flowchart Representation
    doc.add_heading('SYSTEM FLOWCHART LOGIC', level=1)
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    
    steps = [
        "STEP 1: CCTV Live Feed Ingestion",
        "STEP 2: AI inference (Object ID & Count)",
        "STEP 3: Telemetry Sync to Cloud Node",
        "STEP 4: GLOSA Speed Optimization Calculation",
        "STEP 5: Real-time UI Update via WebSockets"
    ]
    
    for i, step in enumerate(steps):
        cell = table.rows[i].cells[0]
        p = cell.paragraphs[0]
        p.text = step
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(12)

    doc.add_page_break()

    # Strategic Value
    doc.add_heading('STRATEGIC ALIGNMENT: AATMANIRBHAR BHARAT', level=1)
    strat_points = [
        "Indigenous Perception: AI trained on Indian traffic (Autos, Bikes) vs foreign LIDAR systems.",
        "Infrastructure Efficiency: Intelligent flow optimization without building new roads.",
        "Sovereign Data: All traffic telemetry handles on national server infrastructure.",
        "Scalability: Designed for Indian municipalities with low-cost hardware compatibility."
    ]
    for pt in strat_points:
        doc.add_paragraph(pt, style='List Bullet')

    doc.save('GLOSA_System_Architecture_Detailed.docx')
    print("Detailed Architecture Doc generated successfully!")

if __name__ == "__main__":
    create_detailed_architecture_doc()
