from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('GLOSA-BHARAT: Project Report & Architecture', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    doc.add_paragraph('Theme: AI for Atmanirbhar Bharat')
    doc.add_paragraph('Concept: Intelligent Traffic Flow & Speed Advisory System')
    doc.add_paragraph('Project Status: Functional Prototype Ready')

    # Section 1: Project File Structure
    doc.add_heading('1. Project File Structure', level=1)
    struct_text = (
        "GLOSA-PROJECT/\n"
        "├── ai-service/                # [PERCEPTION LAYER]\n"
        "│   ├── main.py                # FastAPI server for AI Inference\n"
        "│   ├── requirements.txt       # Python dependencies\n"
        "│   └── models/                # AI Detection Models\n"
        "├── backend/                   # [COORDINATION LAYER]\n"
        "│   ├── server.js              # Express.js Main entry point\n"
        "│   ├── routes/api.js          # Traffic logic & telemetry\n"
        "│   └── package.json\n"
        "└── frontend/                  # [INTERACTION LAYER]\n"
        "    ├── src/components/        # Reusable UI Components\n"
        "    ├── src/pages/             # Main Dashboard UI\n"
        "    └── index.css              # Global Design System"
    )
    p = doc.add_paragraph(struct_text)
    p.style = 'No Spacing'
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

    # Section 2: System Architecture Flowchart
    doc.add_heading('2. System Architecture Flowchart', level=1)
    
    flow_steps = [
        "[1] TRAFFIC JUNCTION: Capture raw CCTV video feed.",
        "[2] AI PERCEPTION LAYER: Identify vehicles/bikes and calculate queue lengths (Python/YOLO).",
        "[3] BACKEND ENGINE: Process telemetry and run GLOSA speed algorithms (Node.js).",
        "[4] DASHBOARD/DRIVER APP: Push real-time speed advisories and GIS map visuals (React/Leaflet).",
        "[5] OUTCOME: Seamless 'Green Wave' flow with zero signal stopping."
    ]
    
    for step in flow_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 3: Technology Stack
    doc.add_heading('3. Technology Stack', level=1)
    tech_table = doc.add_table(rows=1, cols=2)
    hdr_cells = tech_table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technologies'
    
    iters = [
        ('AI & Computer Vision', 'Python, YOLOv8, OpenCV, FastAPI'),
        ('Backend Orchestration', 'Node.js, Express, Axios, WebSockets'),
        ('Frontend Interface', 'React.js, Tailwind CSS, Framer Motion'),
        ('Mapping/GIS', 'React-Leaflet, OpenStreetMap API')
    ]
    
    for layer, tech in iters:
        row_cells = tech_table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech

    # Section 4: Atmanirbhar Bharat Alignment
    doc.add_heading('4. Atmanirbhar Bharat Alignment', level=1)
    points = [
        "Indigenous Software Stack: 100% custom-built for Indian road behavior.",
        "Cost-Effective Innovation: Uses standard CCTV instead of expensive foreign specialized hardware.",
        "Local Data Sovereignty: Keeps traffic telemetry on national servers.",
        "Smart Cities Integration: Ready for deployment in Tier-2/3 Indian cities."
    ]
    for point in points:
        doc.add_paragraph(point, style='List Number')

    # Section 5: Startup Scalability
    doc.add_heading('5. Startup & Business Scalability', level=1)
    doc.add_paragraph(
        "The project is designed as a SaaS platform. Business targets include Government Traffic Police "
        "departments for city-wide monitoring and Logistics fleets for fuel optimization."
    )

    doc.save('GLOSA_Project_Report.docx')
    print("Project Report generated successfully!")

if __name__ == "__main__":
    create_document()
